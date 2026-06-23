"""
Generate a complete set of sample vendor documents for testing the platform.

Produces 5 DOCX files (DOCX extracts cleanly without needing the Tesseract OCR
binary) for a single fictional vendor, one per checklist item:

    1. Business Registration
    2. Tax Certificate
    3. Bank Statement
    4. ID Proof
    5. Audited Financials

Run:  python data/generate_sample_documents.py
Output:  data/sample_vendor/*.docx  (upload these in the UI)
"""
import os

import docx

VENDOR = {
    "name": "TechNova Solutions Ltd",
    "email": "contact@technova.com",
    "business_type": "IT Services",
    "country": "India",
    "address": "4th Floor, Cyber Towers, Hitech City, Hyderabad, Telangana 500081, India",
    "reg_no": "U72200TG2015PTC098765",
    "pan": "AABCT1234J",
    "gstin": "36AABCT1234J1Z5",
}

OUT_DIR = os.path.join(os.path.dirname(__file__), "sample_vendor")


def _doc(title: str) -> docx.Document:
    d = docx.Document()
    d.add_heading(title, level=0)
    return d


def business_registration() -> docx.Document:
    d = _doc("CERTIFICATE OF INCORPORATION")
    d.add_paragraph("Ministry of Corporate Affairs — Government of India")
    d.add_paragraph("Office of the Registrar of Companies, Hyderabad")
    d.add_paragraph("")
    d.add_paragraph(f"Company Name: {VENDOR['name']}")
    d.add_paragraph(f"Corporate Identity Number (CIN): {VENDOR['reg_no']}")
    d.add_paragraph("Date of Incorporation: 14 March 2015")
    d.add_paragraph("Company Type: Private Limited Company")
    d.add_paragraph(f"Registered Office: {VENDOR['address']}")
    d.add_paragraph(
        "This is to certify that the above-named company is incorporated under "
        "the Companies Act, 2013 and that the company is a Private Company "
        "Limited by Shares."
    )
    d.add_paragraph("")
    d.add_paragraph("Registrar of Companies")
    d.add_paragraph("Seal & Signature")
    return d


def tax_certificate() -> docx.Document:
    d = _doc("CERTIFICATE OF TAX REGISTRATION")
    d.add_paragraph("Income Tax Department & GST Authority — Government of India")
    d.add_paragraph("")
    d.add_paragraph(f"Registered Taxpayer: {VENDOR['name']}")
    d.add_paragraph(f"Permanent Account Number (PAN): {VENDOR['pan']}")
    d.add_paragraph(f"GST Identification Number (GSTIN): {VENDOR['gstin']}")
    d.add_paragraph("Tax Registration Status: Active")
    d.add_paragraph("Date of Registration: 02 July 2017")
    d.add_paragraph(f"Principal Place of Business: {VENDOR['address']}")
    d.add_paragraph(
        "This certifies that the above entity is duly registered for income tax "
        "and Goods and Services Tax (GST) purposes and is compliant with filing "
        "obligations as of the date of issue."
    )
    return d


def bank_statement() -> docx.Document:
    d = _doc("BANK ACCOUNT STATEMENT")
    d.add_paragraph("HDFC Bank — Hitech City Branch")
    d.add_paragraph(f"Account Holder: {VENDOR['name']}")
    d.add_paragraph("Account Number: 50200012345678")
    d.add_paragraph("IFSC Code: HDFC0001234")
    d.add_paragraph(f"Address on File: {VENDOR['address']}")
    d.add_paragraph("Statement Period: 01 May 2026 to 31 May 2026")
    d.add_paragraph("")
    table = d.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Description"
    hdr[2].text = "Amount (INR)"
    hdr[3].text = "Balance (INR)"
    rows = [
        ("03 May 2026", "Client payment - Project Atlas", "+1,250,000", "4,820,000"),
        ("11 May 2026", "Vendor payment - Cloud hosting", "-180,000", "4,640,000"),
        ("19 May 2026", "Payroll", "-960,000", "3,680,000"),
        ("28 May 2026", "Client payment - Retainer", "+540,000", "4,220,000"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
    d.add_paragraph("")
    d.add_paragraph("Closing Balance: INR 4,220,000")
    d.add_paragraph("This statement also serves as proof of registered address.")
    return d


def id_proof() -> docx.Document:
    d = _doc("IDENTITY PROOF — AUTHORIZED SIGNATORY")
    d.add_paragraph("Government of India — Aadhaar / Director Identification")
    d.add_paragraph("")
    d.add_paragraph("Name: Priya Ramesh Iyer")
    d.add_paragraph(f"Designation: Managing Director, {VENDOR['name']}")
    d.add_paragraph("Director Identification Number (DIN): 07654321")
    d.add_paragraph("Aadhaar Number: XXXX XXXX 9012")
    d.add_paragraph("Date of Birth: 22 August 1984")
    d.add_paragraph("Nationality: Indian")
    d.add_paragraph(f"Address: {VENDOR['address']}")
    d.add_paragraph(
        "This document serves as government-issued photo identity proof for the "
        "authorized signatory of the vendor entity."
    )
    return d


def audited_financials() -> docx.Document:
    d = _doc("AUDITED FINANCIAL STATEMENTS — FY 2025-26")
    d.add_paragraph(f"Entity: {VENDOR['name']}")
    d.add_paragraph("Auditor: Sharma & Associates, Chartered Accountants")
    d.add_paragraph("Audit Firm Registration No.: 004567S")
    d.add_paragraph("")
    d.add_heading("Statement of Profit & Loss", level=1)
    d.add_paragraph("Total Revenue: INR 250,000,000")
    d.add_paragraph("Total Expenses: INR 198,000,000")
    d.add_paragraph("Profit Before Tax: INR 52,000,000")
    d.add_paragraph("Net Profit After Tax: INR 38,500,000")
    d.add_heading("Balance Sheet", level=1)
    d.add_paragraph("Total Assets: INR 410,000,000")
    d.add_paragraph("Total Liabilities: INR 165,000,000")
    d.add_paragraph("Shareholders' Equity: INR 245,000,000")
    d.add_paragraph("")
    d.add_paragraph(
        "Auditor's Opinion: In our opinion, the financial statements give a true "
        "and fair view of the state of affairs of the company and comply with the "
        "applicable accounting standards."
    )
    d.add_paragraph("Signed: Sharma & Associates, Chartered Accountants")
    return d


def main() -> None:
    os.makedirs(OUT_DIR, exist_ok=True)
    builders = {
        "1_Business_Registration.docx": business_registration,
        "2_Tax_Certificate.docx": tax_certificate,
        "3_Bank_Statement.docx": bank_statement,
        "4_ID_Proof.docx": id_proof,
        "5_Audited_Financials.docx": audited_financials,
    }
    for filename, builder in builders.items():
        path = os.path.join(OUT_DIR, filename)
        builder().save(path)
        print(f"Created {path}")


if __name__ == "__main__":
    main()
