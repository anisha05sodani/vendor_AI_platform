"""Tests for the document loader / text extraction utility."""
from agents.utils.document_loader import (
    load_document,
    load_documents,
    guess_document_type,
)


def test_nonexistent_path_marked_unreadable():
    docs = load_documents(["/no/such/file_xyz.pdf"])
    assert len(docs) == 1
    assert docs[0]["extraction_status"] == "unreadable"


def test_empty_input_returns_empty_list():
    assert load_documents([]) == []


def test_unsupported_extension(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("hello world")
    doc = load_document(str(p))
    assert doc["extraction_status"] == "unsupported"


def test_docx_extraction(tmp_path):
    import docx

    d = docx.Document()
    d.add_heading("Business Registration Certificate", 0)
    d.add_paragraph("Company registration number 12345")
    fp = tmp_path / "reg.docx"
    d.save(str(fp))

    doc = load_document(str(fp))
    assert doc["extraction_status"] == "extracted"
    assert "registration" in doc["extracted_text"].lower()


def test_guess_document_type_from_keywords():
    assert (
        guess_document_type("tax_cert.pdf", "VAT and tax identification number")
        == "tax_identification_document"
    )


def test_guess_document_type_unknown():
    assert guess_document_type("random.pdf", "lorem ipsum dolor") == "unknown"
